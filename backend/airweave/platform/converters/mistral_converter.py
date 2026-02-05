"""Unified converter for documents and images using Mistral OCR."""

import asyncio
import hashlib
import json
import os
import tempfile
import time
from typing import Dict, List, Optional

from httpx import HTTPStatusError, ReadTimeout, TimeoutException
from tenacity import retry, retry_if_exception_type, stop_after_attempt, wait_exponential

from airweave.core.config import settings
from airweave.core.logging import logger
from airweave.platform.converters._base import BaseTextConverter
from airweave.platform.rate_limiters import MistralRateLimiter
from airweave.platform.sync.async_helpers import run_in_thread_pool
from airweave.platform.sync.exceptions import EntityProcessingError, SyncFailureError

# ==================== CONFIGURATION (Module-Level Constants) ====================

# Mistral OCR file size limit
MAX_FILE_SIZE_BYTES = 50_000_000  # 50MB Mistral limit

# Retry configuration (matching Notion pattern)
MAX_RETRIES = 5  # Maximum retry attempts
RETRY_MIN_WAIT = 10  # Minimum wait between retries (seconds)
RETRY_MAX_WAIT = 60  # Maximum wait between retries (seconds)
RETRY_MULTIPLIER = 2  # Exponential backoff multiplier


# ==================== MISTRAL CONVERTER ====================


class MistralConverter(BaseTextConverter):
    """Converts documents and images to markdown using Mistral OCR.

    Supported formats:
    - Documents: PDF, DOCX, PPTX (with splitting)
    - Images: JPG, JPEG, PNG (with compression if >50MB)

    Document splitting (recursive halving):
    - PDF: Split by pages using PyPDF2
    - DOCX: Split by paragraphs using python-docx
    - PPTX: Split by slides using python-pptx
    """

    # Supported formats
    SUPPORTED_DOCUMENT_FORMATS = {".pdf", ".docx", ".pptx"}
    SUPPORTED_IMAGE_FORMATS = {".jpg", ".jpeg", ".png"}
    SUPPORTED_FORMATS = SUPPORTED_DOCUMENT_FORMATS | SUPPORTED_IMAGE_FORMATS

    def __init__(self):
        """Initialize converter with lazy Mistral client initialization."""
        self.max_file_size = MAX_FILE_SIZE_BYTES
        self.rate_limiter = MistralRateLimiter()  # Singleton - shared across all converters in pod
        self._mistral_client = None
        self._mistral_initialized = False

    def _ensure_mistral_client(self):
        """Ensure Mistral client is initialized (lazy initialization).

        Raises:
            SyncFailureError: If Mistral API key not configured or mistralai package not installed
        """
        if self._mistral_initialized:
            return

        # Require Mistral API key
        if not hasattr(settings, "MISTRAL_API_KEY") or not settings.MISTRAL_API_KEY:
            raise SyncFailureError("MISTRAL_API_KEY required for document conversion")

        try:
            from mistralai import Mistral

            self._mistral_client = Mistral(
                api_key=settings.MISTRAL_API_KEY,
                timeout_ms=300_000,  # 5 minute timeout (300 seconds)
            )
            self._mistral_initialized = True
            logger.debug("Mistral client initialized for document conversion")
        except ImportError:
            raise SyncFailureError("mistralai package required but not installed")

    async def _mistral_api_call_with_retry(self, operation):
        """Generic wrapper for Mistral API calls with rate limiting + retry.

        Args:
            operation: Async coroutine to call

        Returns:
            Result from operation

        Raises:
            Exception: If operation fails after retries
        """

        @retry(
            retry=retry_if_exception_type(
                (TimeoutException, ReadTimeout, HTTPStatusError, Exception)
            ),
            stop=stop_after_attempt(MAX_RETRIES),
            wait=wait_exponential(
                multiplier=RETRY_MULTIPLIER, min=RETRY_MIN_WAIT, max=RETRY_MAX_WAIT
            ),
            reraise=True,
        )
        async def _call():
            await self.rate_limiter.acquire()  # Rate limit before each attempt
            return await operation

        return await _call()

    async def convert_batch(self, file_paths: List[str]) -> Dict[str, str]:
        """Convert document files to markdown text using Mistral batch OCR API.

        Args:
            file_paths: List of document file paths to convert

        Returns:
            Dict mapping file_path -> markdown text content (None if failed)
        """
        # Ensure Mistral client is initialized before processing
        self._ensure_mistral_client()

        try:
            # Step 1: Prepare files (split large ones)
            file_chunks_map = await self._prepare_and_split_files(file_paths)

            # Check if all files failed during preparation
            if not file_chunks_map:
                logger.warning("All files failed during preparation")
                return {path: None for path in file_paths}

            # Step 2: Upload all chunks to Mistral
            upload_map = await self._upload_chunks_to_mistral(file_chunks_map)

            # Step 3: Create JSONL batch file with signed URLs
            jsonl_path, custom_id_to_upload_key = await self._create_batch_jsonl(upload_map)

            # Step 4: Submit batch job
            job_id, batch_file_id = await self._submit_batch_job(jsonl_path)

            # Step 5: Poll for completion
            await self._poll_batch_job(job_id, timeout=600)

            # Step 6: Download and parse results
            upload_key_results = await self._download_batch_results(job_id, custom_id_to_upload_key)

            # Step 7: Combine chunks back to original files
            final_results = await self._combine_chunk_results(
                upload_key_results, file_chunks_map, file_paths
            )

            # Step 8: Cleanup
            await self._cleanup_mistral_files(upload_map, batch_file_id)
            await self._cleanup_temp_chunks(file_chunks_map)

            # Step 9: Delete JSONL
            try:
                os.unlink(jsonl_path)
            except Exception:
                pass

            return final_results

        except EntityProcessingError:
            raise
        except Exception as e:
            # Check if this is an infrastructure failure
            error_msg = str(e).lower()
            is_infrastructure_failure = any(
                keyword in error_msg
                for keyword in ["timeout", "api", "network", "connection", "rate limit"]
            )

            if is_infrastructure_failure:
                logger.error(f"Mistral infrastructure failure: {e}")
                raise SyncFailureError(f"Mistral infrastructure failure: {e}")

            # Other unexpected errors
            logger.error(f"Mistral batch conversion failed: {e}")
            raise SyncFailureError(f"Mistral batch conversion failed: {e}")

    # ==================== FILE PREPARATION & SPLITTING ====================

    async def _prepare_and_split_files(self, file_paths: List[str]) -> Dict[str, List[str]]:
        """Check file sizes and split if >50MB.

        Handles errors per-file: if one file fails, others continue processing.

        Args:
            file_paths: Original document file paths

        Returns:
            Dict mapping original_path -> [chunk_paths]
            Failed files are NOT included (will be marked as None in final results)
        """
        file_chunks_map = {}

        # Handle duplicate paths (same file multiple times in batch)
        for idx, path in enumerate(file_paths):
            try:
                file_size = os.path.getsize(path)
                # Create unique key for duplicate paths
                unique_key = f"{path}__batch_idx_{idx}"

                # Get file extension
                _, ext = os.path.splitext(path)
                ext = ext.lower()

                if file_size <= self.max_file_size:
                    # Use as-is
                    file_chunks_map[unique_key] = [path]
                    size_mb = file_size / 1_000_000
                    logger.debug(f"Document {os.path.basename(path)}: {size_mb:.1f}MB (OK)")
                else:
                    # Split into chunks
                    size_mb = file_size / 1_000_000
                    logger.debug(
                        f"Document {os.path.basename(path)} is {size_mb:.1f}MB, splitting..."
                    )
                    chunks = await self._split_large_document(path, file_size, ext)

                    if not chunks:
                        size_mb = file_size / 1_000_000
                        raise EntityProcessingError(
                            f"Failed to split large document: {path} ({size_mb:.1f}MB)"
                        )

                    file_chunks_map[unique_key] = chunks
                    logger.debug(f"Split into {len(chunks)} chunks")

            except EntityProcessingError as e:
                # Entity-level error - log warning and skip this file, continue with others
                logger.warning(f"Skipping file {os.path.basename(path)}: {e}")
                continue
            except Exception as e:
                # Unexpected error - also skip this file
                logger.error(f"Unexpected error preparing {os.path.basename(path)}: {e}")
                continue

        return file_chunks_map

    async def _split_large_document(self, path: str, file_size: int, ext: str) -> List[str]:
        """Dispatch to format-specific handler.

        - Images: Compress if >50MB
        - Documents: Split by pages/paragraphs/slides

        Args:
            path: File path
            file_size: File size in bytes
            ext: File extension (e.g., '.pdf', '.jpg')

        Returns:
            List of chunk/compressed file paths

        Raises:
            EntityProcessingError: If format unsupported or processing fails
        """
        # Images - compress if needed
        if ext in self.SUPPORTED_IMAGE_FORMATS:
            compressed_path = await self._compress_image_if_needed(path, file_size)
            return [compressed_path]

        # Documents - split
        if ext == ".pdf":
            return await self._split_pdf_by_pages(path, file_size)
        elif ext == ".docx":
            return await self._split_docx_by_paragraphs(path, file_size)
        elif ext == ".pptx":
            return await self._split_pptx_by_slides(path, file_size)
        else:
            raise EntityProcessingError(f"Unsupported format: {ext}")

    # ==================== PDF SPLITTING ====================

    async def _split_pdf_by_pages(self, pdf_path: str, file_size: int) -> List[str]:
        """Split PDF into chunks <max_file_size using recursive halving.

        Args:
            pdf_path: Path to large PDF
            file_size: Size in bytes

        Returns:
            List of temporary chunk file paths

        Raises:
            EntityProcessingError: If splitting fails
        """
        try:
            import PyPDF2
        except ImportError:
            raise SyncFailureError("PyPDF2 required to split large PDF but not installed")

        def _split():
            with open(pdf_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                num_pages = len(reader.pages)

                if num_pages == 0:
                    raise EntityProcessingError(f"PDF {pdf_path} has no pages")

                logger.debug(f"PDF has {num_pages} pages, splitting recursively...")

                # Recursively split using halving strategy
                return self._split_pages_recursive(reader, 0, num_pages)

        try:
            return await run_in_thread_pool(_split)
        except EntityProcessingError:
            raise
        except Exception as e:
            raise EntityProcessingError(f"PDF splitting failed for {pdf_path}: {e}")

    def _split_pages_recursive(self, reader, start_idx: int, end_idx: int) -> List[str]:
        """Recursively split page range in half until chunks are small enough.

        Args:
            reader: PyPDF2 PdfReader object
            start_idx: Starting page index
            end_idx: Ending page index (exclusive)

        Returns:
            List of chunk file paths

        Raises:
            EntityProcessingError: If single page exceeds size limit
        """
        import PyPDF2

        num_pages = end_idx - start_idx

        # Create chunk with this page range
        writer = PyPDF2.PdfWriter()
        for i in range(start_idx, end_idx):
            writer.add_page(reader.pages[i])

        # Write to temp file
        temp_fd, temp_path = tempfile.mkstemp(suffix=".pdf")
        with os.fdopen(temp_fd, "wb") as temp_file:
            writer.write(temp_file)

        chunk_size = os.path.getsize(temp_path)

        # Check if chunk is small enough
        if chunk_size <= self.max_file_size:
            logger.debug(
                f"✓ PDF chunk pages {start_idx + 1}-{end_idx}: {chunk_size / 1_000_000:.1f}MB (OK)"
            )
            return [temp_path]

        # Chunk too large - need to split
        os.unlink(temp_path)  # Delete oversized chunk

        # Edge case: single page exceeds limit
        if num_pages == 1:
            limit_mb = self.max_file_size / 1_000_000
            page_mb = chunk_size / 1_000_000
            raise EntityProcessingError(
                f"Single page {start_idx + 1} in PDF ({page_mb:.2f}MB) exceeds limit "
                f"({limit_mb:.2f}MB) - cannot split further"
            )

        # Split in half and recurse
        mid_idx = start_idx + (num_pages // 2)

        chunk_mb = chunk_size / 1_000_000
        logger.debug(
            f"✂️  PDF chunk pages {start_idx + 1}-{end_idx} too large ({chunk_mb:.1f}MB), "
            f"splitting at page {mid_idx}"
        )

        # Recursively process each half
        first_half = self._split_pages_recursive(reader, start_idx, mid_idx)
        second_half = self._split_pages_recursive(reader, mid_idx, end_idx)

        return first_half + second_half

    # ==================== DOCX SPLITTING ====================

    async def _split_docx_by_paragraphs(self, docx_path: str, file_size: int) -> List[str]:
        """Split DOCX by paragraphs using RECURSIVE halving.

        Args:
            docx_path: Path to large DOCX
            file_size: Size in bytes

        Returns:
            List of temporary chunk file paths

        Raises:
            EntityProcessingError: If document has no paragraphs or single paragraph too large
            SyncFailureError: If python-docx package not installed
        """
        try:
            from docx import Document
        except ImportError:
            raise SyncFailureError("python-docx package required but not installed")

        def _split():
            doc = Document(docx_path)
            all_paragraphs = list(doc.paragraphs)

            if not all_paragraphs:
                raise EntityProcessingError(f"DOCX {docx_path} has no paragraphs")

            logger.debug(f"DOCX has {len(all_paragraphs)} paragraphs, splitting recursively...")

            # Recursively split paragraph range
            return self._split_paragraphs_recursive(all_paragraphs, 0, len(all_paragraphs))

        try:
            return await run_in_thread_pool(_split)
        except EntityProcessingError:
            raise
        except Exception as e:
            raise EntityProcessingError(f"DOCX splitting failed for {docx_path}: {e}")

    def _split_paragraphs_recursive(self, paragraphs: List, start: int, end: int) -> List[str]:
        """Recursively split paragraph range until chunk <50MB.

        Args:
            paragraphs: List of paragraph objects
            start: Starting paragraph index
            end: Ending paragraph index (exclusive)

        Returns:
            List of chunk file paths

        Raises:
            EntityProcessingError: If single paragraph exceeds size limit
        """
        from docx import Document

        # Create chunk with this range
        chunk_doc = Document()
        for i in range(start, end):
            chunk_doc.add_paragraph(paragraphs[i].text, style=paragraphs[i].style)

        # Save and check size
        temp_fd, temp_path = tempfile.mkstemp(suffix=".docx")
        os.close(temp_fd)
        chunk_doc.save(temp_path)
        chunk_size = os.path.getsize(temp_path)

        if chunk_size <= self.max_file_size:
            logger.debug(
                f"✓ DOCX chunk paragraphs {start + 1}-{end}: {chunk_size / 1_000_000:.1f}MB (OK)"
            )
            return [temp_path]

        # Chunk too large - split
        os.unlink(temp_path)

        # Edge case: Single paragraph too large
        if end - start == 1:
            limit_mb = self.max_file_size / 1_000_000
            chunk_mb = chunk_size / 1_000_000
            raise EntityProcessingError(
                f"Single paragraph in DOCX ({chunk_mb:.2f}MB) exceeds limit "
                f"({limit_mb:.2f}MB) - cannot split further"
            )

        # Split in half
        mid = start + (end - start) // 2
        logger.debug(f"✂️  DOCX chunk paragraphs {start + 1}-{end} too large, splitting at {mid}")

        first_half = self._split_paragraphs_recursive(paragraphs, start, mid)
        second_half = self._split_paragraphs_recursive(paragraphs, mid, end)

        return first_half + second_half

    # ==================== PPTX SPLITTING ====================

    async def _split_pptx_by_slides(self, pptx_path: str, file_size: int) -> List[str]:
        """Split PPTX by slides using RECURSIVE halving.

        Args:
            pptx_path: Path to large PPTX
            file_size: Size in bytes

        Returns:
            List of temporary chunk file paths

        Raises:
            EntityProcessingError: If presentation has no slides or single slide too large
            SyncFailureError: If python-pptx package not installed
        """
        try:
            from pptx import Presentation
        except ImportError:
            raise SyncFailureError("python-pptx package required but not installed")

        def _split():
            prs = Presentation(pptx_path)
            num_slides = len(prs.slides)

            if num_slides == 0:
                raise EntityProcessingError(f"PPTX {pptx_path} has no slides")

            logger.debug(f"PPTX has {num_slides} slides, splitting recursively...")

            # Recursively split slide range
            return self._split_slides_recursive(pptx_path, 0, num_slides)

        try:
            return await run_in_thread_pool(_split)
        except EntityProcessingError:
            raise
        except Exception as e:
            raise EntityProcessingError(f"PPTX splitting failed for {pptx_path}: {e}")

    def _split_slides_recursive(self, template_path: str, start: int, end: int) -> List[str]:
        """Recursively split slide range until chunk <50MB.

        Args:
            template_path: Path to original PPTX (used as template)
            start: Starting slide index
            end: Ending slide index (exclusive)

        Returns:
            List of chunk file paths

        Raises:
            EntityProcessingError: If single slide exceeds size limit
        """
        import shutil

        # For simplicity, copy the original and use as chunk
        # (Full slide extraction is complex; Mistral OCR handles visual content)
        temp_fd, temp_path = tempfile.mkstemp(suffix=".pptx")
        os.close(temp_fd)
        shutil.copy2(template_path, temp_path)

        chunk_size = os.path.getsize(temp_path)

        if chunk_size <= self.max_file_size:
            logger.debug(
                f"✓ PPTX chunk slides {start + 1}-{end}: {chunk_size / 1_000_000:.1f}MB (OK)"
            )
            return [temp_path]

        # Too large - split
        os.unlink(temp_path)

        # Edge case: Single slide too large
        if end - start == 1:
            limit_mb = self.max_file_size / 1_000_000
            slide_mb = chunk_size / 1_000_000
            raise EntityProcessingError(
                f"Single slide in PPTX ({slide_mb:.2f}MB) exceeds limit "
                f"({limit_mb:.2f}MB) - cannot split further"
            )

        # Split in half
        mid = start + (end - start) // 2
        logger.debug(f"✂️  PPTX chunk slides {start + 1}-{end} too large, splitting at {mid}")

        first_half = self._split_slides_recursive(template_path, start, mid)
        second_half = self._split_slides_recursive(template_path, mid, end)

        return first_half + second_half

    # ==================== IMAGE COMPRESSION ====================

    async def _compress_image_if_needed(self, image_path: str, file_size: int) -> str:
        """Compress image if >50MB using Pillow.

        Args:
            image_path: Path to image file
            file_size: File size in bytes

        Returns:
            Path to compressed image (or original if <50MB)

        Raises:
            SyncFailureError: If Pillow not installed
            EntityProcessingError: If cannot compress below 50MB
        """
        if file_size <= self.max_file_size:
            return image_path  # No compression needed

        try:
            from PIL import Image
        except ImportError:
            raise SyncFailureError(
                "Pillow package required for image compression but not installed"
            )

        def _compress():
            img = Image.open(image_path)
            _, ext = os.path.splitext(image_path)
            ext = ext.lower()

            # PNG compression: Convert to JPEG (lossy but effective)
            if ext == ".png":
                # Convert RGBA to RGB if needed
                if img.mode in ("RGBA", "LA", "P"):
                    background = Image.new("RGB", img.size, (255, 255, 255))
                    background.paste(img, mask=img.split()[-1] if img.mode == "RGBA" else None)
                    img = background

                # Try quality levels for JPEG conversion
                for quality in range(85, 19, -10):
                    temp_fd, temp_path = tempfile.mkstemp(suffix=".jpg")
                    os.close(temp_fd)

                    img.save(temp_path, "JPEG", quality=quality, optimize=True)
                    compressed_size = os.path.getsize(temp_path)

                    if compressed_size <= self.max_file_size:
                        size_mb = file_size / 1_000_000
                        compressed_mb = compressed_size / 1_000_000
                        logger.debug(
                            f"Compressed PNG→JPEG {os.path.basename(image_path)}: "
                            f"{size_mb:.2f}MB → {compressed_mb:.2f}MB (quality={quality}%)"
                        )
                        return temp_path

                    os.unlink(temp_path)

            # JPEG compression: Use quality parameter
            else:
                for quality in range(85, 19, -10):
                    temp_fd, temp_path = tempfile.mkstemp(suffix=ext)
                    os.close(temp_fd)

                    img.save(temp_path, quality=quality, optimize=True)
                    compressed_size = os.path.getsize(temp_path)

                    if compressed_size <= self.max_file_size:
                        size_mb = file_size / 1_000_000
                        compressed_mb = compressed_size / 1_000_000
                        logger.debug(
                            f"Compressed {os.path.basename(image_path)}: "
                            f"{size_mb:.2f}MB → {compressed_mb:.2f}MB (quality={quality}%)"
                        )
                        return temp_path

                    os.unlink(temp_path)

            # Cannot compress enough
            size_mb = file_size / 1_000_000
            limit_mb = self.max_file_size / 1_000_000
            raise EntityProcessingError(
                f"Cannot compress {os.path.basename(image_path)} below {limit_mb:.2f}MB "
                f"(original: {size_mb:.2f}MB)"
            )

        return await run_in_thread_pool(_compress)

    # ==================== MISTRAL API CALLS (WITH RATE LIMITING) ====================

    async def _upload_chunks_to_mistral(
        self, file_chunks_map: Dict[str, List[str]]
    ) -> Dict[str, dict]:
        """Upload all chunk files to Mistral and get signed URLs.

        Args:
            file_chunks_map: Mapping of original_path -> [chunk_paths]

        Returns:
            Dict mapping upload_key -> {'file_id', 'signed_url'}
        """
        # Flatten all chunks with their unique keys
        all_chunks = []
        for unique_key, chunks in file_chunks_map.items():
            for chunk_path in chunks:
                all_chunks.append((unique_key, chunk_path))

        logger.debug(f"Uploading {len(all_chunks)} document chunks to Mistral...")

        # Upload concurrently with semaphore (10 concurrent max)
        semaphore = asyncio.Semaphore(10)
        upload_map = {}

        failed_uploads = []  # Track failed uploads

        async def _upload_one(unique_key: str, chunk_path: str):
            async with semaphore:
                try:
                    # Upload file with retry + rate limiting
                    with open(chunk_path, "rb") as f:
                        file_response = await self._mistral_api_call_with_retry(
                            self._mistral_client.files.upload_async(
                                file={"file_name": os.path.basename(chunk_path), "content": f},
                                purpose="ocr",
                            )
                        )
                    file_id = file_response.id

                    # Get signed URL with retry + rate limiting
                    signed_url_response = await self._mistral_api_call_with_retry(
                        self._mistral_client.files.get_signed_url_async(file_id=file_id)
                    )
                    signed_url = signed_url_response.url

                    # Store result
                    upload_key = f"{unique_key}__chunk_{chunk_path}"
                    upload_map[upload_key] = {"file_id": file_id, "signed_url": signed_url}

                    logger.debug(f"Uploaded {os.path.basename(chunk_path)}")

                except Exception as e:
                    # Log error but don't fail entire batch - track for later
                    logger.error(f"Upload failed for {os.path.basename(chunk_path)}: {e}")
                    upload_key = f"{unique_key}__chunk_{chunk_path}"
                    failed_uploads.append((upload_key, chunk_path, str(e)))

        await asyncio.gather(
            *[_upload_one(uk, cp) for uk, cp in all_chunks], return_exceptions=True
        )

        logger.debug(f"Uploaded {len(upload_map)} chunks successfully")
        return upload_map

    async def _create_batch_jsonl(self, upload_map: Dict[str, dict]) -> tuple:
        """Create JSONL batch file with signed URLs.

        Detects file type (image vs document) and uses appropriate Mistral API type.

        Args:
            upload_map: Dict of upload_key -> {'file_id', 'signed_url'}

        Returns:
            Tuple of (jsonl_path, custom_id_to_upload_key_mapping)
        """

        def _create():
            temp_fd, jsonl_path = tempfile.mkstemp(suffix=".jsonl")
            custom_id_to_upload_key = {}

            with os.fdopen(temp_fd, "w") as f:
                for upload_key, upload_info in upload_map.items():
                    # Use hash as custom_id (avoid path length issues)
                    custom_id = hashlib.md5(upload_key.encode()).hexdigest()
                    custom_id_to_upload_key[custom_id] = upload_key

                    # Extract file path and extension to determine type
                    file_path = upload_key.split("__chunk_")[1]
                    _, ext = os.path.splitext(file_path)
                    ext = ext.lower()

                    # Choose document_url or image_url based on extension
                    if ext in self.SUPPORTED_IMAGE_FORMATS:
                        doc_type = "image_url"
                    else:
                        doc_type = "document_url"

                    entry = {
                        "custom_id": custom_id,
                        "body": {
                            "document": {
                                "type": doc_type,
                                doc_type: upload_info["signed_url"],
                            }
                        },
                    }
                    f.write(json.dumps(entry) + "\n")

            logger.debug(f"Created batch JSONL with {len(upload_map)} entries: {jsonl_path}")
            return jsonl_path, custom_id_to_upload_key

        return await run_in_thread_pool(_create)

    async def _submit_batch_job(self, jsonl_path: str) -> tuple:
        """Upload JSONL and create batch job.

        Args:
            jsonl_path: Path to JSONL file

        Returns:
            Tuple of (job_id, batch_file_id)

        Raises:
            EntityProcessingError: If submission fails
        """
        try:
            # Upload JSONL with retry + rate limiting
            with open(jsonl_path, "rb") as f:
                batch_file_response = await self._mistral_api_call_with_retry(
                    self._mistral_client.files.upload_async(
                        file={"file_name": "batch.jsonl", "content": f}, purpose="batch"
                    )
                )
            batch_file_id = batch_file_response.id

            # Create batch job with retry + rate limiting
            job_response = await self._mistral_api_call_with_retry(
                self._mistral_client.batch.jobs.create_async(
                    input_files=[batch_file_id], model="mistral-ocr-latest", endpoint="/v1/ocr"
                )
            )
            job_id = job_response.id

            logger.debug(f"Submitted batch job {job_id} (batch file: {batch_file_id})")
            return job_id, batch_file_id

        except Exception as e:
            raise EntityProcessingError(f"Failed to submit Mistral batch job: {e}")

    async def _poll_batch_job(self, job_id: str, timeout: int = 600):
        """Poll batch job until complete.

        Args:
            job_id: Mistral batch job ID
            timeout: Maximum wait time in seconds

        Raises:
            EntityProcessingError: If job fails or times out
        """
        start_time = time.time()
        last_log = 0

        while True:
            elapsed = time.time() - start_time

            if elapsed > timeout:
                raise EntityProcessingError(f"Mistral batch job {job_id} timeout after {timeout}s")

            # Check status with retry + rate limiting
            try:
                job = await self._mistral_api_call_with_retry(
                    lambda: self._mistral_client.batch.jobs.get(job_id=job_id)
                )
            except Exception as e:
                raise EntityProcessingError(f"Failed to check batch job status: {e}")

            if job.status == "SUCCESS":
                logger.debug(f"Batch job {job_id} completed successfully")
                return job

            elif job.status == "FAILED":
                error_msg = getattr(job, "error", "Unknown error")
                raise EntityProcessingError(f"Mistral batch job {job_id} failed: {error_msg}")

            # Log progress every 10 seconds
            if elapsed - last_log >= 10:
                succeeded = getattr(job, "succeeded_requests", 0)
                total = getattr(job, "total_requests", 0)
                logger.debug(f"Batch OCR progress: {succeeded}/{total} ({int(elapsed)}s elapsed)")
                last_log = elapsed

            await asyncio.sleep(2)  # Poll every 2 seconds

    async def _download_batch_results(  # noqa: C901
        self, job_id: str, custom_id_to_upload_key: Dict[str, str]
    ) -> Dict[str, str]:
        """Download and parse batch results.

        Args:
            job_id: Mistral batch job ID
            custom_id_to_upload_key: Mapping of custom_id -> upload_key

        Returns:
            Dict mapping upload_key -> markdown content
        """
        try:
            # Get job with retry + rate limiting to get output file ID
            job = await self._mistral_api_call_with_retry(
                lambda: self._mistral_client.batch.jobs.get(job_id=job_id)
            )
            if not hasattr(job, "output_file") or not job.output_file:
                raise Exception("Batch job has no output file")

            # Download output file with retry + rate limiting
            response = await self._mistral_api_call_with_retry(
                lambda: self._mistral_client.files.download(file_id=job.output_file)
            )

            # Handle different response types
            try:
                if isinstance(response, bytes):
                    result_data = response
                elif hasattr(response, "read"):
                    result_data = response.read()
                elif hasattr(response, "content"):
                    result_data = response.content
                elif hasattr(response, "text"):
                    result_data = response.text.encode("utf-8")
                else:
                    result_data = str(response).encode("utf-8")
            except Exception:
                result_data = response

        except Exception as e:
            raise EntityProcessingError(f"Failed to download batch results: {e}")

        # Parse JSONL results
        upload_key_results = {}

        # Handle both bytes and str
        if isinstance(result_data, bytes):
            result_text = result_data.decode("utf-8")
        else:
            result_text = str(result_data)

        for line in result_text.strip().split("\n"):
            if not line:
                continue

            try:
                item = json.loads(line)
                custom_id = item.get("custom_id")

                if not custom_id or custom_id not in custom_id_to_upload_key:
                    logger.warning(f"Unknown custom_id in results: {custom_id}")
                    continue

                upload_key = custom_id_to_upload_key[custom_id]

                # Check if Mistral returned an error for this specific file
                if "error" in item and item["error"]:
                    error_info = item["error"]
                    if isinstance(error_info, dict):
                        error_msg = error_info.get("message", "Unknown error")
                        error_code = error_info.get("code", "N/A")
                    else:
                        error_msg = str(error_info)
                        error_code = "N/A"
                    logger.error(f"Mistral OCR failed for {upload_key}: [{error_code}] {error_msg}")
                    upload_key_results[upload_key] = None
                    continue

                # Extract markdown from response
                if "response" in item and "body" in item["response"]:
                    pages = item["response"]["body"].get("pages", [])
                    if pages:
                        markdown = "\n\n".join([p.get("markdown", "") for p in pages])
                        upload_key_results[upload_key] = markdown
                    else:
                        logger.warning(f"No pages in OCR result for {upload_key}")
                        upload_key_results[upload_key] = None
                else:
                    logger.warning(f"No OCR result for {upload_key}")
                    upload_key_results[upload_key] = None

            except Exception as e:
                logger.error(f"Failed to parse batch result line: {e}")
                continue

        logger.debug(f"Parsed {len(upload_key_results)} OCR results from batch")
        return upload_key_results

    @staticmethod
    def _warn_if_empty_markdown(
        markdown: Optional[str],
        file_name: str,
        chunk_idx: Optional[int] = None,
        total_chunks: Optional[int] = None,
    ):
        if isinstance(markdown, str) and not markdown.strip():
            if chunk_idx is None:
                logger.warning(f"OCR returned empty markdown for {file_name}")
            else:
                logger.warning(
                    f"OCR returned empty markdown for chunk {chunk_idx}/{total_chunks} "
                    f"of {file_name}"
                )

    def _process_single_chunk_result(
        self,
        upload_key_results: Dict[str, str],
        unique_key: str,
        chunk_path: str,
        file_name: str,
    ) -> Optional[str]:
        upload_key = f"{unique_key}__chunk_{chunk_path}"
        markdown = upload_key_results.get(upload_key)

        if markdown is None:
            logger.error(
                f"Conversion failed for {file_name}: Upload or OCR failed "
                f"(check logs above for details)"
            )
            return None

        self._warn_if_empty_markdown(markdown, file_name)
        return markdown

    def _process_multi_chunk_result(
        self,
        upload_key_results: Dict[str, str],
        unique_key: str,
        chunk_paths: List[str],
        file_name: str,
    ) -> Optional[str]:
        combined = ""

        for idx, chunk_path in enumerate(chunk_paths, start=1):
            upload_key = f"{unique_key}__chunk_{chunk_path}"
            chunk_md = upload_key_results.get(upload_key)

            if chunk_md is None:
                logger.error(
                    f"Conversion failed for {file_name}: Chunk {idx}/{len(chunk_paths)} "
                    f"upload or OCR failed"
                )
                return None

            self._warn_if_empty_markdown(chunk_md, file_name, idx, len(chunk_paths))

            if idx > 1:
                combined += "\n\n---\n\n"
            combined += chunk_md

        logger.debug(f"Successfully combined {len(chunk_paths)} chunks for {file_name}")
        return combined

    async def _combine_chunk_results(
        self,
        upload_key_results: Dict[str, str],
        file_chunks_map: Dict[str, List[str]],
        original_file_paths: List[str],
    ) -> Dict[str, str]:  # noqa: C901
        """Combine chunk markdown back to original files.

        Args:
            upload_key_results: Dict of upload_key -> markdown (None if OCR failed)
            file_chunks_map: Dict of unique_key -> [chunk_paths]
            original_file_paths: Original list of file paths requested

        Returns:
            Dict mapping original_path -> combined markdown (None if failed)

        Note: None results are EXPLICITLY logged with reason:
        - File skipped during preparation (e.g., page too large)
        - Upload failed
        - Mistral OCR failed
        - Chunk recombination failed
        """
        final_results = {}

        # Process files that were successfully prepared
        for unique_key, chunk_paths in file_chunks_map.items():
            # Strip __batch_idx_X suffix to get original path
            original_path = unique_key.split("__batch_idx_")[0]
            file_name = os.path.basename(original_path)

            if len(chunk_paths) == 1:
                final_results[original_path] = self._process_single_chunk_result(
                    upload_key_results, unique_key, chunk_paths[0], file_name
                )
            else:
                final_results[original_path] = self._process_multi_chunk_result(
                    upload_key_results, unique_key, chunk_paths, file_name
                )

        # Mark files that failed during preparation as None
        for path in original_file_paths:
            if path not in final_results:
                # Explicit: File was skipped during preparation
                logger.error(
                    f"Conversion failed for {os.path.basename(path)}: "
                    f"File skipped during preparation (e.g., page/paragraph/slide too large)"
                )
                final_results[path] = None

        return final_results

    # ==================== CLEANUP ====================

    async def _cleanup_mistral_files(self, upload_map: Dict[str, dict], batch_file_id: str):
        """Delete uploaded files from Mistral cloud (best effort).

        Args:
            upload_map: Dict of upload_key -> {'file_id', 'signed_url'}
            batch_file_id: ID of uploaded batch JSONL file
        """

        async def _delete(file_id):
            try:
                # Delete with retry + rate limiting (best effort)
                await self._mistral_api_call_with_retry(
                    self._mistral_client.files.delete_async(file_id=file_id)
                )
            except Exception:
                pass  # Best effort cleanup - don't fail if deletion fails

        # Collect all file IDs
        file_ids = [info["file_id"] for info in upload_map.values()]
        file_ids.append(batch_file_id)

        await asyncio.gather(*[_delete(fid) for fid in file_ids], return_exceptions=True)
        logger.debug(f"Cleaned up {len(file_ids)} files from Mistral")

    async def _cleanup_temp_chunks(self, file_chunks_map: Dict[str, List[str]]):
        """Delete temporary chunk files (best effort).

        Args:
            file_chunks_map: Dict of unique_key -> [chunk_paths]
        """
        for unique_key, chunk_paths in file_chunks_map.items():
            original_path = unique_key.split("__batch_idx_")[0]
            for chunk_path in chunk_paths:
                if chunk_path != original_path:  # Don't delete original
                    try:
                        os.unlink(chunk_path)
                        logger.debug(f"Deleted temp chunk {chunk_path}")
                    except Exception:
                        pass  # Best effort
