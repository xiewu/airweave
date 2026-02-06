"""Recursive-halving document splitters for oversized files.

Each splitter loads a document, counts its atomic units (pages, paragraphs),
then recursively halves the range until every chunk is below the size limit.

The generic algorithm lives in :class:`RecursiveSplitter`; format-specific
implementations only override three hooks: ``load``, ``unit_count``, and
``write_range``.
"""

from __future__ import annotations

import asyncio
import os
import tempfile
from abc import ABC, abstractmethod
from typing import Any

import aiofiles.os

from airweave.core.logging import logger
from airweave.platform.sync.exceptions import EntityProcessingError, SyncFailureError


class RecursiveSplitter(ABC):
    """Base class for recursive-halving document splitters.

    Subclasses implement three hooks:

    * ``load``        -- open / parse the document (runs in a thread).
    * ``unit_count``  -- return the number of splittable units.
    * ``write_range`` -- write units ``[start, end)`` to a temp file (runs in a thread).
    """

    # Human-readable name for log messages (e.g. "PDF", "DOCX").
    format_label: str = "document"
    # Unit name for log messages (e.g. "page", "paragraph").
    unit_label: str = "unit"

    # ---- hooks for subclasses ------------------------------------------------

    @abstractmethod
    async def load(self, path: str) -> Any:
        """Parse the source document. Called once per ``split()`` invocation."""

    @abstractmethod
    def unit_count(self, source: Any) -> int:
        """Return the total number of splittable units in *source*."""

    @abstractmethod
    async def write_range(self, source: Any, start: int, end: int) -> str:
        """Write units ``[start, end)`` to a new temp file.

        Returns:
            Path to the temporary file.
        """

    # ---- public API ----------------------------------------------------------

    async def split(self, path: str, max_bytes: int) -> list[str]:
        """Split *path* into chunks each at most *max_bytes*.

        Returns:
            Ordered list of temp-file paths.

        Raises:
            EntityProcessingError: If a single unit exceeds *max_bytes*.
        """
        source = await self.load(path)
        count = self.unit_count(source)

        if count == 0:
            raise EntityProcessingError(f"{self.format_label} {path} has no {self.unit_label}s")

        logger.debug(
            f"{self.format_label} has {count} {self.unit_label}s, splitting recursively..."
        )
        return await self._halve(source, 0, count, max_bytes)

    # ---- recursive halving ---------------------------------------------------

    async def _halve(self, source: Any, start: int, end: int, max_bytes: int) -> list[str]:
        temp_path = await self.write_range(source, start, end)
        stat = await aiofiles.os.stat(temp_path)
        chunk_size = stat.st_size

        if chunk_size <= max_bytes:
            size_mb = chunk_size / 1_000_000
            logger.debug(
                f"{self.format_label} chunk {self.unit_label}s "
                f"{start + 1}-{end}: {size_mb:.1f}MB (OK)"
            )
            return [temp_path]

        # Chunk too large -- delete and recurse.
        await aiofiles.os.remove(temp_path)

        if end - start == 1:
            limit_mb = max_bytes / 1_000_000
            chunk_mb = chunk_size / 1_000_000
            raise EntityProcessingError(
                f"Single {self.unit_label} {start + 1} in {self.format_label} "
                f"({chunk_mb:.2f}MB) exceeds limit ({limit_mb:.2f}MB) -- cannot split further"
            )

        mid = start + (end - start) // 2
        chunk_mb = chunk_size / 1_000_000
        logger.debug(
            f"{self.format_label} chunk {self.unit_label}s {start + 1}-{end} too large "
            f"({chunk_mb:.1f}MB), splitting at {self.unit_label} {mid}"
        )

        first_half = await self._halve(source, start, mid, max_bytes)
        second_half = await self._halve(source, mid, end, max_bytes)
        return first_half + second_half


# ===========================================================================
# PDF splitter
# ===========================================================================


class PdfSplitter(RecursiveSplitter):
    """Split a PDF by page ranges using PyPDF2."""

    format_label = "PDF"
    unit_label = "page"

    async def load(self, path: str) -> Any:
        """Load a PDF and return a ``PdfReader`` with materialised pages."""
        try:
            import PyPDF2
        except ImportError:
            raise SyncFailureError("PyPDF2 required to split large PDFs but not installed")

        def _load():
            with open(path, "rb") as fh:
                reader = PyPDF2.PdfReader(fh)
                # Materialize pages so the file handle can close.
                # PdfReader keeps pages in memory after read.
                _ = len(reader.pages)
                return reader

        return await asyncio.to_thread(_load)

    def unit_count(self, source: Any) -> int:
        """Return the number of pages in the PDF."""
        return len(source.pages)

    async def write_range(self, source: Any, start: int, end: int) -> str:
        """Write a page range to a temporary PDF file."""
        import PyPDF2

        def _write():
            writer = PyPDF2.PdfWriter()
            for i in range(start, end):
                writer.add_page(source.pages[i])

            fd, temp_path = tempfile.mkstemp(suffix=".pdf")
            with os.fdopen(fd, "wb") as fh:
                writer.write(fh)
            return temp_path

        return await asyncio.to_thread(_write)


# ===========================================================================
# DOCX splitter
# ===========================================================================


class DocxSplitter(RecursiveSplitter):
    """Split a DOCX by paragraph ranges using python-docx."""

    format_label = "DOCX"
    unit_label = "paragraph"

    async def load(self, path: str) -> Any:
        """Load a DOCX and return a list of paragraphs."""
        try:
            from docx import Document
        except ImportError:
            raise SyncFailureError("python-docx package required but not installed")

        def _load():
            doc = Document(path)
            return list(doc.paragraphs)

        return await asyncio.to_thread(_load)

    def unit_count(self, source: Any) -> int:
        """Return the number of paragraphs."""
        return len(source)

    async def write_range(self, source: Any, start: int, end: int) -> str:
        """Write a paragraph range to a temporary DOCX file."""
        from docx import Document

        def _write():
            chunk_doc = Document()
            for i in range(start, end):
                chunk_doc.add_paragraph(source[i].text, style=source[i].style)

            fd, temp_path = tempfile.mkstemp(suffix=".docx")
            os.close(fd)
            chunk_doc.save(temp_path)
            return temp_path

        return await asyncio.to_thread(_write)
